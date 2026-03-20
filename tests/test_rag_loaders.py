import os
import tempfile

from src.rag.loaders import load_pdf, load_csv, load_xlsx, load_docx, load_all_from_dir


def test_load_csv():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False, encoding="utf-8") as f:
        f.write("nome,preco,modelo\n")
        f.write("Filtro de Oleo,25.90,CB300\n")
        f.write("Pastilha de Freio,45.00,CG160\n")
        path = f.name

    try:
        docs = load_csv(path)
        assert len(docs) >= 1
        assert any("Filtro" in d.page_content for d in docs)
        assert docs[0].metadata["source"] == path
    finally:
        os.unlink(path)


def test_load_xlsx():
    import openpyxl

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["nome", "preco", "modelo"])
    ws.append(["Filtro de Oleo", 25.90, "CB300"])
    wb.save(path)

    try:
        docs = load_xlsx(path)
        assert len(docs) >= 1
        assert any("Filtro" in d.page_content for d in docs)
    finally:
        os.unlink(path)


def test_load_docx():
    from docx import Document as DocxDocument

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    doc = DocxDocument()
    doc.add_paragraph("Catalogo de pecas Temporalis AI")
    doc.add_paragraph("Filtro de Oleo CB300 - R$ 25.90")
    doc.save(path)

    try:
        docs = load_docx(path)
        assert len(docs) >= 1
        assert any("Filtro" in d.page_content for d in docs)
    finally:
        os.unlink(path)


def test_load_pdf_extracts_text():
    """load_pdf should handle a blank PDF (no extractable text)."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        writer.write(f)
        path = f.name
    try:
        docs = load_pdf(path)
        assert isinstance(docs, list)
        # Blank page has no text, so docs should be empty
        assert len(docs) == 0
    finally:
        os.unlink(path)


def test_load_docx_empty_returns_empty():
    """load_docx with no paragraphs returns empty list."""
    from docx import Document as DocxDocument

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        doc = DocxDocument()
        doc.save(path)
        result = load_docx(path)
        assert result == []
    finally:
        os.unlink(path)


def test_load_all_from_dir():
    with tempfile.TemporaryDirectory() as tmpdir:
        csv_path = os.path.join(tmpdir, "products.csv")
        with open(csv_path, "w", encoding="utf-8") as f:
            f.write("nome,preco\nFiltro,25.90\n")

        txt_path = os.path.join(tmpdir, "notes.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("ignored file")

        docs = load_all_from_dir(tmpdir)
        assert len(docs) >= 1
        assert all(hasattr(d, "page_content") for d in docs)
