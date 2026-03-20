import os

import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".csv", ".xlsx", ".docx"}


def load_pdf(path: str) -> list[Document]:
    reader = PdfReader(path)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text.strip(),
                metadata={"source": path, "page": i + 1, "file_type": "pdf"},
            ))
    return docs


def load_csv(path: str) -> list[Document]:
    df = pd.read_csv(path, dtype=str).fillna("")
    docs = []
    for _, row in df.iterrows():
        content = " | ".join(f"{col}: {val}" for col, val in row.items() if val)
        if content.strip():
            docs.append(Document(
                page_content=content,
                metadata={"source": path, "file_type": "csv"},
            ))
    return docs


def load_xlsx(path: str) -> list[Document]:
    df = pd.read_excel(path, dtype=str, engine="openpyxl").fillna("")
    docs = []
    for _, row in df.iterrows():
        content = " | ".join(f"{col}: {val}" for col, val in row.items() if val)
        if content.strip():
            docs.append(Document(
                page_content=content,
                metadata={"source": path, "file_type": "xlsx"},
            ))
    return docs


def load_docx(path: str) -> list[Document]:
    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return []
    full_text = "\n".join(paragraphs)
    return [Document(
        page_content=full_text,
        metadata={"source": path, "file_type": "docx"},
    )]


LOADER_MAP = {
    ".pdf": load_pdf,
    ".csv": load_csv,
    ".xlsx": load_xlsx,
    ".docx": load_docx,
}


def load_all_from_dir(directory: str) -> list[Document]:
    docs = []
    for filename in os.listdir(directory):
        ext = os.path.splitext(filename)[1].lower()
        if ext in LOADER_MAP:
            filepath = os.path.join(directory, filename)
            loader = LOADER_MAP[ext]
            docs.extend(loader(filepath))
    return docs
